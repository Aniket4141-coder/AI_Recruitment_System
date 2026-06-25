from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Sequence

import pandas as pd
import torch
from docx import Document

from ranking_diagnostics import diagnose_top_candidates
from src.behavioral_scorer import BehavioralScorer
from src.experience_scorer import ExperienceScorer
from src.parser import Candidate
from src.product_fit import ProductFitScorer
from src.ranking_engine import RankingEngine
from src.reasoning_generator import ReasoningGenerator
from src.semantic_matcher import SemanticMatcher
from src.skill_matcher import SkillMatcher
from src.title_scorer import TitleScorer

DEFAULT_CANDIDATES_PATH = Path("data/candidates.jsonl")
DEFAULT_JD_PATH = Path("data/job_description.docx")
DEFAULT_REPORT_PATH = Path("validation_report.json")
DEFAULT_TOP20_CSV = Path("validation_top20.csv")

OUTPUT_COLUMNS = [
    "candidate_id",
    "rank",
    "final_score",
    "semantic_score",
    "skill_score",
    "behavior_score",
    "experience_score",
    "product_fit_score",
    "title_score",
    "reasoning",
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Validate ranking quality for an AI recruitment JD.")
    parser.add_argument("--input", default=str(DEFAULT_CANDIDATES_PATH), help="Path to candidates.jsonl.")
    parser.add_argument("--jd", default=str(DEFAULT_JD_PATH), help="Path to the job description DOCX.")
    parser.add_argument("--limit", type=int, default=None, help="Optional cap for quick validation runs.")
    parser.add_argument("--top-n", type=int, default=20, help="How many top candidates to diagnose.")
    parser.add_argument("--chunk-size", type=int, default=512, help="Batch size for semantic scoring.")
    parser.add_argument("--device", choices=["auto", "cpu", "cuda"], default="auto", help="Execution device.")
    parser.add_argument("--report", default=str(DEFAULT_REPORT_PATH), help="Path to write JSON validation report.")
    parser.add_argument("--top-csv", default=str(DEFAULT_TOP20_CSV), help="Path to write the Top N CSV.")
    return parser.parse_args()


def resolve_device(mode: str) -> str:
    if mode == "cpu":
        return "cpu"
    if mode == "cuda":
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA was requested but is not available on this machine.")
        return "cuda"
    return "cuda" if torch.cuda.is_available() else "cpu"


def iter_docx_text(document: Document) -> Iterator[str]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text
    for table in document.tables:
        for row in table.rows:
            text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if text:
                yield text


def load_job_description(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"Job description not found: {path}")
    document = Document(str(path))
    text = "\n".join(iter_docx_text(document)).strip()
    if not text:
        raise ValueError(f"Job description file is empty: {path}")
    return text


def iter_candidates(path: Path, limit: int | None = None) -> Iterator[Dict[str, Any]]:
    if not path.exists():
        raise FileNotFoundError(f"Candidate file not found: {path}")
    with path.open("r", encoding="utf-8") as handle:
        for line_number, line in enumerate(handle, start=1):
            if limit is not None and line_number > limit:
                break
            line = line.strip()
            if not line:
                continue
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(record, dict):
                yield record


def chunked(iterable: Iterable[Dict[str, Any]], chunk_size: int) -> Iterator[List[Dict[str, Any]]]:
    chunk: List[Dict[str, Any]] = []
    for item in iterable:
        chunk.append(item)
        if len(chunk) >= max(1, chunk_size):
            yield chunk
            chunk = []
    if chunk:
        yield chunk


def initialize_scorers(device: str) -> Dict[str, Any]:
    semantic_matcher = SemanticMatcher()
    if semantic_matcher.model is not None:
        semantic_matcher.model = semantic_matcher.model.to(device)
        semantic_matcher.model.eval()

    skill_matcher = SkillMatcher()
    return {
        "semantic": semantic_matcher,
        "skill": skill_matcher,
        "experience": ExperienceScorer(),
        "behavioral": BehavioralScorer(),
        "product_fit": ProductFitScorer(),
        "title": TitleScorer(),
        "ranking": RankingEngine(),
        "reasoning": ReasoningGenerator(),
    }


def score_candidates(
    raw_candidates: Sequence[Dict[str, Any]],
    jd_text: str,
    scorers: Dict[str, Any],
) -> List[Dict[str, Any]]:
    candidates = [Candidate(record) for record in raw_candidates]
    candidate_texts = [candidate.get_candidate_text() for candidate in candidates]
    semantic_scores = scorers["semantic"].calculate_similarity_batch(jd_text, candidate_texts)

    rows: List[Dict[str, Any]] = []
    for candidate, semantic_score in zip(candidates, semantic_scores):
        skill_score = float(scorers["skill"].calculate_skill_score(candidate.skills))
        experience_score = float(scorers["experience"].calculate(candidate.experience))
        behavior_score = float(scorers["behavioral"].calculate(candidate.redrob_signals))
        product_fit_score = float(scorers["product_fit"].calculate(candidate))
        title_score = float(scorers["title"].calculate(candidate.current_title, candidate.career_history))
        final_score = float(
            scorers["ranking"].calculate_final_score(
                semantic_score,
                skill_score,
                experience_score,
                behavior_score,
                product_fit_score,
                title_score,
            )
        )
        explanation = scorers["reasoning"].generate(candidate, final_score)

        rows.append(
            {
                "candidate_id": candidate.id or "",
                "current_title": candidate.current_title,
                "experience": float(candidate.experience or 0),
                "skills": candidate.skills,
                "career_history": candidate.career_history,
                "career_history_text": candidate.get_career_text(),
                "semantic_score": round(semantic_score, 4),
                "skill_score": skill_score,
                "behavior_score": behavior_score,
                "experience_score": experience_score,
                "product_fit_score": product_fit_score,
                "title_score": title_score,
                "final_score": final_score,
                "reasoning": explanation,
            }
        )

    return rows


def rank_rows(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows.sort(key=lambda row: (-row["final_score"], row["candidate_id"]))
    for rank, row in enumerate(rows, start=1):
        row["rank"] = rank
    return rows


def build_top20_dataframe(rows: Sequence[Dict[str, Any]]) -> pd.DataFrame:
    df = pd.DataFrame(rows[:20])
    if df.empty:
        return pd.DataFrame(columns=OUTPUT_COLUMNS)
    df = df[OUTPUT_COLUMNS + ["current_title", "experience"]].copy()
    return df


def print_report(report: Dict[str, Any]) -> None:
    print("\nRANKING VALIDATION REPORT")
    print("=" * 30)
    print(f"Top candidates analyzed: {report['top_count']}")
    print(f"Obvious false positives: {report['irrelevant_count']}")
    print(f"False positive rate: {report['false_positive_rate']:.2%}")

    print("\nCandidate Diagnostics")
    for item in report["candidate_diagnostics"]:
        label = "FALSE POSITIVE" if item["is_irrelevant"] else "OK"
        print(f"- Rank {item['rank']}: {item['candidate_id']} | {item['current_title']} | {label}")
        for reason in item["reasons"][:3]:
            print(f"  - {reason}")

    print("\nSuggested Fixes")
    for fix in report["suggested_fixes"]:
        print(f"- {fix}")


def main() -> int:
    args = parse_args()
    device = resolve_device(args.device)
    jd_text = load_job_description(Path(args.jd))
    scorers = initialize_scorers(device)

    rows: List[Dict[str, Any]] = []
    for chunk in chunked(iter_candidates(Path(args.input), limit=args.limit), args.chunk_size):
        rows.extend(score_candidates(chunk, jd_text, scorers))

    if not rows:
        raise RuntimeError("No valid candidates were found to validate.")

    rank_rows(rows)
    top_rows = rows[: max(1, args.top_n)]
    report = diagnose_top_candidates(top_rows)
    report["job_description_chars"] = len(jd_text)
    report["total_ranked_candidates"] = len(rows)

    top_df = build_top20_dataframe(rows[: max(1, args.top_n)])
    top_df.to_csv(args.top_csv, index=False)

    with Path(args.report).open("w", encoding="utf-8") as handle:
        json.dump(report, handle, indent=2)

    print_report(report)
    print(f"\nSaved report to: {args.report}")
    print(f"Saved top-{args.top_n} CSV to: {args.top_csv}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
