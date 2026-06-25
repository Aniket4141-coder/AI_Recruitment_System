from __future__ import annotations

import argparse
import csv
import json
import logging
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence

import torch
from docx import Document
from tqdm import tqdm

from src.behavioral_scorer import BehavioralScorer
from src.experience_scorer import ExperienceScorer
from src.parser import Candidate
from src.product_fit import ProductFitScorer
from src.ranking_engine import RankingEngine
from src.explainability import ExplainabilityGenerator
from src.semantic_matcher import SemanticMatcher
from src.skill_matcher import SkillMatcher
from src.title_scorer import TitleScorer

OUTPUT_COLUMNS = [
    "candidate_id",
    "rank",
    "final_score",
    "semantic_score",
    "skill_score",
    "behavior_score",
    "experience_score",
    "product_fit_score",
    "title_score",
    "reasoning",
]

LOGGER = logging.getLogger("recruitment_pipeline")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Production-ready AI recruitment ranking pipeline."
    )
    parser.add_argument(
        "--input",
        default="data/candidates.jsonl",
        help="Path to the candidates JSONL file.",
    )
    parser.add_argument(
        "--jd",
        default="data/job_description.docx",
        help="Path to the job description DOCX file.",
    )
    parser.add_argument(
        "--output",
        default="submission.csv",
        help="Path to write the final ranking CSV.",
    )
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=512,
        help="Number of candidates to process per semantic batch.",
    )
    parser.add_argument(
        "--limit",
        type=int,
        default=None,
        help="Optional cap for quick validation runs.",
    )
    parser.add_argument(
        "--device",
        choices=["auto", "cpu", "cuda"],
        default="auto",
        help="Execution device for semantic embeddings.",
    )
    parser.add_argument(
        "--log-level",
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        help="Logging verbosity.",
    )
    return parser.parse_args()


def configure_logging(level: str) -> None:
    logging.basicConfig(
        level=getattr(logging, level),
        format="%(asctime)s | %(levelname)s | %(message)s",
    )


def resolve_device(mode: str) -> str:
    if mode == "cpu":
        return "cpu"
    if mode == "cuda":
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA was requested but is not available on this machine.")
        return "cuda"
    return "cuda" if torch.cuda.is_available() else "cpu"


def iter_docx_text(document: Document) -> Iterator[str]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in document.tables:
        for row in table.rows:
            cell_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if cell_text:
                yield cell_text


def load_job_description(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(f"Job description not found: {path}")

    document = Document(str(path))
    text_parts = list(iter_docx_text(document))
    jd_text = "\n".join(text_parts).strip()

    if not jd_text:
        raise ValueError(f"Job description file is empty: {path}")

    return jd_text


def iter_candidates(path: Path, limit: Optional[int] = None) -> Iterator[Dict[str, Any]]:
    if not path.exists():
        raise FileNotFoundError(f"Candidate file not found: {path}")

    with path.open("r", encoding="utf-8") as handle:
        for line_number, line in enumerate(handle, start=1):
            if limit is not None and line_number > limit:
                break

            line = line.strip()
            if not line:
                continue

            try:
                record = json.loads(line)
            except json.JSONDecodeError as exc:
                LOGGER.warning("Skipping invalid JSON at line %s: %s", line_number, exc)
                continue

            if not isinstance(record, dict):
                LOGGER.warning("Skipping non-object record at line %s", line_number)
                continue

            yield record


def chunked(iterable: Iterable[Dict[str, Any]], chunk_size: int) -> Iterator[List[Dict[str, Any]]]:
    chunk: List[Dict[str, Any]] = []
    for item in iterable:
        chunk.append(item)
        if len(chunk) >= chunk_size:
            yield chunk
            chunk = []
    if chunk:
        yield chunk


def initialize_scorers(device: str) -> Dict[str, Any]:
    semantic_matcher = SemanticMatcher()
    if semantic_matcher.model is not None:
        semantic_matcher.model = semantic_matcher.model.to(device)
        semantic_matcher.model.eval()

    skill_matcher = SkillMatcher()

    return {
        "semantic": semantic_matcher,
        "skill": skill_matcher,
        "experience": ExperienceScorer(),
        "behavioral": BehavioralScorer(),
        "product_fit": ProductFitScorer(),
        "title": TitleScorer(),
        "ranking": RankingEngine(),
        "explainability": ExplainabilityGenerator(skill_matcher=skill_matcher),
    }


def score_chunk(
    raw_candidates: Sequence[Dict[str, Any]],
    jd_text: str,
    scorers: Dict[str, Any],
) -> List[Dict[str, Any]]:
    candidates = [Candidate(record) for record in raw_candidates]
    candidate_texts = [candidate.get_candidate_text() for candidate in candidates]
    semantic_scores = scorers["semantic"].calculate_similarity_batch(jd_text, candidate_texts)

    rows: List[Dict[str, Any]] = []
    for candidate, semantic_score in zip(candidates, semantic_scores):
        skill_score = scorers["skill"].calculate_skill_score(candidate.skills)
        experience_score = scorers["experience"].calculate(candidate.experience)
        behavior_score = scorers["behavioral"].calculate(candidate.redrob_signals)
        product_fit_score = scorers["product_fit"].calculate(candidate)
        title_score = scorers["title"].calculate(candidate.current_title, candidate.career_history)

        final_score = scorers["ranking"].calculate_final_score(
            semantic_score,
            skill_score,
            experience_score,
            behavior_score,
            product_fit_score,
            title_score,
        )

        explanation = scorers["explainability"].generate(
            candidate=candidate,
            jd_text=jd_text,
            semantic_score=semantic_score,
            skill_score=skill_score,
            behavior_score=behavior_score,
            experience_score=experience_score,
            product_fit_score=product_fit_score,
            title_score=title_score,
        )

        rows.append(
            {
                "candidate_id": candidate.id or "",
                "final_score": final_score,
                "semantic_score": round(float(semantic_score), 4),
                "skill_score": skill_score,
                "behavior_score": behavior_score,
                "experience_score": experience_score,
                "product_fit_score": product_fit_score,
                "title_score": title_score,
                "reasoning": (
                    "Why selected:\n"
                    f"{explanation['why_selected']}\n\n"
                    "Why not selected:\n"
                    f"{explanation['why_not_selected']}"
                ),
            }
        )

    return rows


def rank_rows(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows.sort(key=lambda row: (-row["final_score"], row["candidate_id"]))

    for rank, row in enumerate(rows, start=1):
        row["rank"] = rank

    return rows


def write_submission(rows: Sequence[Dict[str, Any]], output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    with output_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=OUTPUT_COLUMNS)
        writer.writeheader()

        for row in rows:
            writer.writerow(
                {
                    "candidate_id": row["candidate_id"],
                    "rank": row["rank"],
                    "final_score": f"{row['final_score']:.4f}",
                    "semantic_score": f"{row['semantic_score']:.4f}",
                    "skill_score": f"{row['skill_score']:.4f}",
                    "behavior_score": f"{row['behavior_score']:.4f}",
                    "experience_score": f"{row['experience_score']:.4f}",
                    "product_fit_score": f"{row['product_fit_score']:.4f}",
                    "title_score": f"{row['title_score']:.4f}",
                    "reasoning": row["reasoning"],
                }
            )


def main() -> int:
    args = parse_args()
    configure_logging(args.log_level)

    input_path = Path(args.input)
    jd_path = Path(args.jd)
    output_path = Path(args.output)

    LOGGER.info("Loading job description from %s", jd_path)
    jd_text = load_job_description(jd_path)
    LOGGER.info("Job description loaded (%s characters)", len(jd_text))

    device = resolve_device(args.device)
    LOGGER.info("Using device: %s", device)

    scorers = initialize_scorers(device)

    LOGGER.info("Reading candidates from %s", input_path)
    candidate_stream = iter_candidates(input_path, limit=args.limit)

    total_scored = 0
    ranked_rows: List[Dict[str, Any]] = []

    for chunk in tqdm(
        chunked(candidate_stream, max(1, args.chunk_size)),
        desc="Scoring chunks",
        unit="chunk",
    ):
        scored_rows = score_chunk(chunk, jd_text, scorers)
        ranked_rows.extend(scored_rows)
        total_scored += len(scored_rows)

    if not ranked_rows:
        raise RuntimeError("No valid candidate records were found to score.")

    LOGGER.info("Scored %s candidates", total_scored)
    LOGGER.info("Ranking candidates")
    rank_rows(ranked_rows)

    LOGGER.info("Writing submission to %s", output_path)
    write_submission(ranked_rows, output_path)

    LOGGER.info("Done. Wrote %s ranked candidates", len(ranked_rows))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())