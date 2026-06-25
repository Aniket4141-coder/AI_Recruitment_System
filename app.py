from __future__ import annotations

import html
import json
import os
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Sequence, Tuple

import pandas as pd
import plotly.express as px
import streamlit as st
import torch
from docx import Document

from src.behavioral_scorer import BehavioralScorer
from src.experience_scorer import ExperienceScorer
from src.parser import Candidate
from src.product_fit import ProductFitScorer
from src.ranking_engine import RankingEngine
from src.reasoning_generator import ReasoningGenerator
from src.semantic_matcher import SemanticMatcher
from src.skill_matcher import SkillMatcher
from src.title_scorer import TitleScorer

APP_TITLE = "AI Recruitment System"
APP_SUBTITLE = "Semantic resume matching, candidate ranking, and recruiter-friendly explainability."

DEFAULT_CANDIDATES_PATH = Path("data/candidates.jsonl")
DEFAULT_JD_PATH = Path("data/job_description.docx")

OUTPUT_COLUMNS = [
    "candidate_id",
    "rank",
    "final_score",
    "semantic_score",
    "skill_score",
    "behavior_score",
    "experience_score",
    "product_fit_score",
    "title_score",
    "reasoning",
]


st.set_page_config(
    page_title=APP_TITLE,
    page_icon="??",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>
    :root {
        --bg: #f5f7fb;
        --card: #ffffff;
        --ink: #0f172a;
        --muted: #64748b;
        --line: #e2e8f0;
        --accent: #1d4ed8;
        --accent-2: #0f766e;
    }

    .stApp {
        background: radial-gradient(circle at top left, rgba(29, 78, 216, 0.08), transparent 26%),
                    radial-gradient(circle at top right, rgba(15, 118, 110, 0.08), transparent 24%),
                    linear-gradient(180deg, #f8fafc 0%, #f5f7fb 100%);
    }

    h1, h2, h3, h4 {
        color: var(--ink);
        letter-spacing: -0.02em;
    }

    .hero {
        padding: 1.25rem 1.4rem;
        border: 1px solid var(--line);
        border-radius: 20px;
        background: rgba(255, 255, 255, 0.9);
        box-shadow: 0 16px 40px rgba(15, 23, 42, 0.06);
        margin-bottom: 1rem;
    }

    .hero h1 {
        margin: 0;
        font-size: 2.35rem;
        font-weight: 800;
        background: linear-gradient(135deg, #0f172a 0%, #1d4ed8 52%, #0f766e 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }

    .hero p {
        margin-top: 0.35rem;
        color: var(--muted);
        font-size: 1.02rem;
    }

    .section-card {
        border: 1px solid var(--line);
        border-radius: 18px;
        background: rgba(255, 255, 255, 0.95);
        box-shadow: 0 12px 32px rgba(15, 23, 42, 0.05);
        padding: 1rem 1rem 0.25rem 1rem;
        margin-bottom: 1rem;
    }

    .section-title {
        font-size: 1.05rem;
        font-weight: 700;
        color: var(--ink);
        margin-bottom: 0.85rem;
    }

    .subtle-note {
        color: var(--muted);
        font-size: 0.92rem;
    }

    div[data-testid="stMetricValue"] {
        font-size: 1.55rem;
        font-weight: 800;
        color: var(--accent);
    }

    .reason-box {
        border: 1px solid var(--line);
        border-radius: 14px;
        padding: 0.9rem 1rem;
        background: #f8fafc;
        white-space: pre-wrap;
    }
</style>
""",
    unsafe_allow_html=True,
)


def render_hero() -> None:
    st.markdown(
        f"""
<div class="hero">
    <h1>{APP_TITLE}</h1>
    <p>{APP_SUBTITLE}</p>
</div>
""",
        unsafe_allow_html=True,
    )


def read_docx(file_obj) -> str:
    document = Document(file_obj)
    parts: List[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def read_txt(file_obj) -> str:
    if hasattr(file_obj, "getvalue"):
        return file_obj.getvalue().decode("utf-8", errors="replace").strip()
    return file_obj.read().decode("utf-8", errors="replace").strip()


@st.cache_data(show_spinner=True)
def load_candidates(file_path: str, file_mtime: float) -> List[Dict[str, Any]]:
    if not os.path.exists(file_path):
        return []

    records: List[Dict[str, Any]] = []
    with open(file_path, "r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if not line:
                continue
            try:
                record = json.loads(line)
            except json.JSONDecodeError:
                continue
            if isinstance(record, dict):
                records.append(record)
    return records


@st.cache_resource(show_spinner=True)
def load_scorers(device_mode: str) -> Dict[str, Any]:
    semantic_matcher = SemanticMatcher()
    if semantic_matcher.model is not None:
        semantic_matcher.model = semantic_matcher.model.to(resolve_device(device_mode))
        semantic_matcher.model.eval()

    skill_matcher = SkillMatcher()
    return {
        "semantic": semantic_matcher,
        "skill": skill_matcher,
        "experience": ExperienceScorer(),
        "behavioral": BehavioralScorer(),
        "product_fit": ProductFitScorer(),
        "title": TitleScorer(),
        "ranking": RankingEngine(),
        "reasoning": ReasoningGenerator(),
    }


def resolve_device(device_mode: str) -> str:
    if device_mode == "cpu":
        return "cpu"
    if device_mode == "cuda":
        if not torch.cuda.is_available():
            raise RuntimeError("CUDA was selected but is not available on this machine.")
        return "cuda"
    return "cuda" if torch.cuda.is_available() else "cpu"


def load_job_description_text(uploaded_file) -> Tuple[str, str]:
    if uploaded_file is not None:
        suffix = (uploaded_file.name.split(".")[-1] or "").lower()
        if suffix == "docx":
            return read_docx(uploaded_file), uploaded_file.name
        if suffix == "txt":
            return read_txt(uploaded_file), uploaded_file.name
        return "", uploaded_file.name

    if DEFAULT_JD_PATH.exists():
        with DEFAULT_JD_PATH.open("rb") as handle:
            return read_docx(handle), str(DEFAULT_JD_PATH)

    return "", ""


def pre_filter_candidates(
    candidates: Sequence[Dict[str, Any]],
    title_query: str,
    min_exp: int,
    max_exp: int,
    country_query: str,
    open_to_work: bool,
    max_notice: int,
) -> List[Dict[str, Any]]:
    filtered: List[Dict[str, Any]] = []
    title_query = (title_query or "").strip().lower()
    country_query = (country_query or "").strip().lower()

    for record in candidates:
        profile = record.get("profile", {}) if isinstance(record, dict) else {}
        signals = record.get("redrob_signals", {}) if isinstance(record, dict) else {}

        years = profile.get("years_of_experience", 0) or 0
        if years < min_exp or years > max_exp:
            continue

        current_title = (profile.get("current_title") or "").lower()
        if title_query and title_query not in current_title:
            continue

        country = (profile.get("country") or "").lower()
        if country_query and country_query not in country:
            continue

        if open_to_work and not signals.get("open_to_work_flag", False):
            continue

        notice_days = signals.get("notice_period_days", 180) or 180
        if notice_days > max_notice:
            continue

        filtered.append(record)

    return filtered


def chunked(items: Sequence[Dict[str, Any]], chunk_size: int) -> Iterator[List[Dict[str, Any]]]:
    step = max(1, chunk_size)
    for start in range(0, len(items), step):
        yield list(items[start : start + step])


def score_candidates(
    candidate_records: Sequence[Dict[str, Any]],
    jd_text: str,
    scorers: Dict[str, Any],
    chunk_size: int,
    progress_bar,
    status_box,
) -> Tuple[List[Dict[str, Any]], Dict[str, Candidate]]:
    scored_rows: List[Dict[str, Any]] = []
    candidate_lookup: Dict[str, Candidate] = {}
    total = len(candidate_records)

    if total == 0:
        return scored_rows, candidate_lookup

    processed = 0
    for chunk in chunked(candidate_records, chunk_size):
        candidates = [Candidate(record) for record in chunk]
        candidate_texts = [candidate.get_candidate_text() for candidate in candidates]
        semantic_scores = scorers["semantic"].calculate_similarity_batch(jd_text, candidate_texts)

        for candidate, semantic_score in zip(candidates, semantic_scores):
            candidate_id = candidate.id or f"candidate_{processed + 1:05d}"
            candidate_lookup[candidate_id] = candidate

            skill_score = float(scorers["skill"].calculate_skill_score(candidate.skills))
            experience_score = float(scorers["experience"].calculate(candidate.experience))
            behavior_score = float(scorers["behavioral"].calculate(candidate.redrob_signals))
            product_fit_score = float(scorers["product_fit"].calculate(candidate))
            title_score = float(scorers["title"].calculate(candidate.current_title, candidate.career_history))
            final_score = float(
                scorers["ranking"].calculate_final_score(
                    semantic_score,
                    skill_score,
                    experience_score,
                    behavior_score,
                    product_fit_score,
                    title_score,
                )
            )
            reasoning = scorers["reasoning"].generate(candidate, final_score)

            scored_rows.append(
                {
                    "candidate_id": candidate_id,
                    "candidate": candidate,
                    "summary": candidate.summary,
                    "current_title": candidate.current_title,
                    "experience": float(candidate.experience or 0),
                    "skills": candidate.skills,
                    "semantic_score": round(float(semantic_score), 4),
                    "skill_score": round(skill_score, 4),
                    "behavior_score": round(behavior_score, 4),
                    "experience_score": round(experience_score, 4),
                    "product_fit_score": round(product_fit_score, 4),
                    "title_score": round(title_score, 4),
                    "final_score": round(final_score, 4),
                    "reasoning": reasoning,
                    "profile": candidate,
                }
            )

            processed += 1
            if progress_bar is not None:
                progress_bar.progress(min(processed / max(1, total), 0.98))
            if status_box is not None:
                status_box.write(f"Scored {processed:,} of {total:,} candidates...")

    return scored_rows, candidate_lookup


def rank_rows(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows.sort(key=lambda row: (-row["final_score"], row["candidate_id"]))
    for rank, row in enumerate(rows, start=1):
        row["rank"] = rank
    return rows


def build_submission_dataframe(rows: Sequence[Dict[str, Any]]) -> pd.DataFrame:

    rows = rows[:100]

    if not rows:
        return pd.DataFrame(columns=OUTPUT_COLUMNS)

    df = pd.DataFrame(rows)
    df = df[OUTPUT_COLUMNS].copy()
    df["rank"] = df["rank"].astype(int)
    for column in [
        "final_score",
        "semantic_score",
        "skill_score",
        "behavior_score",
        "experience_score",
        "product_fit_score",
        "title_score",
    ]:
        df[column] = pd.to_numeric(df[column], errors="coerce").fillna(0.0)
    return df


def build_score_distribution_figure(rows: Sequence[Dict[str, Any]]):
    if not rows:
        return None

    df = pd.DataFrame(rows)
    fig = px.histogram(
        df,
        x="final_score",
        nbins=24,
        title="Final Score Distribution",
        color_discrete_sequence=["#2563eb"],
    )
    fig.update_layout(
        margin=dict(l=10, r=10, t=50, b=10),
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(248,250,252,0.95)",
        height=420,
        xaxis_title="Final Score",
        yaxis_title="Candidates",
    )
    return fig


def render_candidate_card(row: Dict[str, Any]) -> None:
    candidate = row["candidate"]
    reason_text = html.escape(str(row["reasoning"]))
    summary_text = html.escape(str(candidate.summary or "No summary available."))
    skills_text = ", ".join(candidate.skills) if candidate.skills else "No structured skills found."
    skills_text = html.escape(skills_text)

    st.markdown(
        f"""
<div class="subtle-note"><strong>Rank #{row['rank']}</strong> | Candidate ID: {row['candidate_id']} | Current Title: {html.escape(candidate.current_title or 'N/A')} | Experience: {float(candidate.experience or 0):.1f} years | Final Score: {float(row['final_score']):.4f}</div>
""",
        unsafe_allow_html=True,
    )
    st.markdown(f"""**Candidate Summary**  
{summary_text}""")
    st.markdown(f"""**Skills**  
{skills_text}""")

    score_cols = st.columns(4)
    score_cols[0].metric("Semantic Score", f"{row['semantic_score']:.4f}")
    score_cols[1].metric("Skill Score", f"{row['skill_score']:.4f}")
    score_cols[2].metric("Experience Score", f"{row['experience_score']:.4f}")
    score_cols[3].metric("Behavioral Score", f"{row['behavior_score']:.4f}")

    more_cols = st.columns(3)
    more_cols[0].metric("Product Fit", f"{row['product_fit_score']:.4f}")
    more_cols[1].metric("Title Score", f"{row['title_score']:.4f}")
    more_cols[2].metric("Final Score", f"{row['final_score']:.4f}")

    st.markdown("**Reasoning**")
    st.markdown(f'<div class="reason-box">{reason_text}</div>', unsafe_allow_html=True)


def render_top20_table(top20: pd.DataFrame) -> None:
    table_df = top20[["rank", "candidate_id", "current_title", "experience", "final_score"]].copy()
    table_df = table_df.rename(
        columns={
            "rank": "Rank",
            "candidate_id": "Candidate ID",
            "current_title": "Current Title",
            "experience": "Experience",
            "final_score": "Final Score",
        }
    )
    table_df["Experience"] = table_df["Experience"].astype(float).round(1)
    table_df["Final Score"] = table_df["Final Score"].astype(float).round(4)

    st.dataframe(table_df, use_container_width=True, hide_index=True)


def render_download_button(submission_df: pd.DataFrame) -> None:
    csv_bytes = submission_df.to_csv(index=False).encode("utf-8")
    st.download_button(
        label="Download CSV",
        data=csv_bytes,
        file_name="submission.csv",
        mime="text/csv",
        use_container_width=True,
    )


def main() -> None:
    render_hero()

    with st.sidebar:
        st.markdown("### Filters")
        title_filter = st.text_input("Role title filter", placeholder="e.g. machine learning")
        country_filter = st.text_input("Country filter", placeholder="e.g. India")
        exp_range = st.slider("Years of experience", min_value=0, max_value=30, value=(3, 15))
        open_to_work_only = st.checkbox("Open to work only", value=False)
        notice_period_max = st.slider("Max notice period (days)", min_value=0, max_value=180, value=90)

        st.markdown("---")
        st.markdown("### Performance")
        pool_size = st.slider("Candidate pool size", min_value=100, max_value=10000, value=1000, step=100)
        chunk_size = st.slider("Semantic batch size", min_value=32, max_value=1024, value=256, step=32)
        device_mode = st.selectbox("Compute device", ["auto", "cpu", "cuda"], index=0)

    uploaded_jd = st.file_uploader("Upload Job Description (.docx or .txt)", type=["docx", "txt"])
    jd_text, jd_source_name = load_job_description_text(uploaded_jd)

    if not jd_text:
        st.warning("Upload a job description or keep `data/job_description.docx` in the repository.")
        st.stop()

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Job Description</div>', unsafe_allow_html=True)
    st.text_area("Loaded Job Description", jd_text, height=240, disabled=True, label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

    run_clicked = st.button("?? Find Best Candidates", type="primary", use_container_width=True)

    if run_clicked:
        candidates_path = DEFAULT_CANDIDATES_PATH
        candidates_mtime = candidates_path.stat().st_mtime if candidates_path.exists() else 0.0
        candidates_db = load_candidates(str(candidates_path), candidates_mtime)

        if not candidates_db:
            st.error(f"No candidates were found at `{candidates_path}`.")
            st.stop()

        filtered_candidates = pre_filter_candidates(
            candidates_db,
            title_filter,
            exp_range[0],
            exp_range[1],
            country_filter,
            open_to_work_only,
            notice_period_max,
        )

        available_count = len(filtered_candidates)
        eval_count = min(available_count, pool_size)

        metric_cols = st.columns(4)
        metric_cols[0].metric("Available Matches", f"{available_count:,}")
        metric_cols[1].metric("Selected for Ranking", f"{eval_count:,}")
        metric_cols[2].metric("Pool Cap", f"{pool_size:,}")
        metric_cols[3].metric("Batch Size", f"{chunk_size:,}")

        if eval_count == 0:
            st.info("No candidates matched the current filters. Broaden the filters and try again.")
            st.session_state.pop("ranking_rows", None)
            st.session_state.pop("submission_df", None)
            st.stop()

        if eval_count > 5000 and device_mode != "cuda":
            st.warning("Ranking more than 5,000 candidates on CPU may take some time. Consider tighter filters or GPU mode.")

        scorers = load_scorers(device_mode)
        progress = st.progress(0.0)
        status = st.empty()
        with st.spinner("Running ranking pipeline..."):
            eval_pool = filtered_candidates[:eval_count]
            scored_rows, candidate_lookup = score_candidates(
                eval_pool,
                jd_text,
                scorers,
                chunk_size,
                progress,
                status,
            )
            ranked_rows = rank_rows(scored_rows)
            submission_df = build_submission_dataframe(ranked_rows)

        st.session_state["ranking_rows"] = ranked_rows
        st.session_state["candidate_lookup"] = candidate_lookup
        st.session_state["submission_df"] = submission_df
        progress.progress(1.0)
        status.success(f"Ranking completed for {len(ranked_rows):,} candidates.")

    if "ranking_rows" not in st.session_state:
        st.info("Click ?? Find Best Candidates to rank the candidate pool and generate the shortlist.")
        st.stop()

    ranked_rows = st.session_state["ranking_rows"]
    submission_df = st.session_state["submission_df"]

    if not ranked_rows:
        st.warning("No ranked candidates are available.")
        st.stop()

    top20 = pd.DataFrame(ranked_rows).head(20).copy()
    top20_table = top20[["rank", "candidate_id", "current_title", "experience", "final_score"]].copy()

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Top 20 Candidates</div>', unsafe_allow_html=True)
    render_top20_table(top20_table)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Download</div>', unsafe_allow_html=True)
    render_download_button(submission_df)
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Score Distribution</div>', unsafe_allow_html=True)
    score_fig = build_score_distribution_figure(ranked_rows)
    if score_fig is not None:
        st.plotly_chart(score_fig, use_container_width=True)
    else:
        st.info("No score distribution data is available yet.")
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown('<div class="section-title">Candidate Cards</div>', unsafe_allow_html=True)
    for row in ranked_rows[:20]:
        candidate = row["candidate"]
        with st.expander(
            f"Rank #{row['rank']} | {row['candidate_id']} | {candidate.current_title or 'Unknown title'} | Final Score: {row['final_score']:.4f}",
            expanded=(row["rank"] == 1),
        ):
            render_candidate_card(row)
    st.markdown('</div>', unsafe_allow_html=True)


if __name__ == "__main__":
    main()
